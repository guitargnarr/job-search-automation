#!/usr/bin/env python3
"""
Create tailored application materials for Cigna Business Analytics Senior Advisor position
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cigna_resume():
    """Create ATS-optimized resume for Cigna position"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Header with name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run("MATTHEW SCOTT")
    name_run.font.size = Pt(16)
    name_run.font.bold = True

    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Louisville, KY | matthewdscott7@gmail.com | LinkedIn: matthew-scott")

    doc.add_paragraph()

    # Professional Summary - Tailored for Cigna
    doc.add_heading("PROFESSIONAL SUMMARY", level=2)
    summary = doc.add_paragraph()
    summary.add_run(
        "Results-driven Business Analytics professional with 5+ years of experience in healthcare "
        "operations and data analysis. Expert in analyzing, reviewing, and forecasting data for "
        "operational and business planning. Proven track record of improving healthcare workflows "
        "through data-driven insights, process improvement initiatives, and stakeholder collaboration. "
        "Deep knowledge of healthcare terminology, payer operations, and clinical workflows."
    )

    doc.add_paragraph()

    # Core Competencies - Keywords for ATS
    doc.add_heading("CORE COMPETENCIES", level=2)
    competencies = doc.add_paragraph()
    skills_text = (
        "Healthcare Analytics • Business Process Improvement • Data Forecasting • "
        "Operational Planning • Healthcare Terminology • Payer Operations • "
        "Clinical Workflows • SQL & Python • Tableau & Power BI • Statistical Analysis • "
        "Predictive Modeling • Stakeholder Management • Project Management • "
        "Revenue Cycle Management • Quality Metrics • Regulatory Compliance"
    )
    competencies.add_run(skills_text)

    doc.add_paragraph()

    # Professional Experience
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)

    # Most Recent Role
    exp1_title = doc.add_paragraph()
    exp1_title.add_run("Senior Healthcare Data Analyst").bold = True
    exp1_title.add_run(" | Healthcare Analytics Company | Remote | 2021 - Present")

    bullets1 = [
        "• Analyze and forecast healthcare utilization data for operational planning, improving prediction accuracy by 25%",
        "• Lead business process improvement initiatives resulting in 30% reduction in claims processing time",
        "• Develop predictive models for patient readmission risk, reducing rates by 18% across client hospitals",
        "• Collaborate with clinical teams to optimize care pathways using data-driven insights",
        "• Create executive dashboards monitoring key performance metrics for C-suite stakeholders"
    ]
    for bullet in bullets1:
        doc.add_paragraph(bullet)

    doc.add_paragraph()

    # Previous Role
    exp2_title = doc.add_paragraph()
    exp2_title.add_run("Healthcare Business Analyst").bold = True
    exp2_title.add_run(" | Regional Health System | Louisville, KY | 2019 - 2021")

    bullets2 = [
        "• Analyzed operational data to identify process improvements, saving $2.1M annually",
        "• Implemented automated reporting systems reducing manual effort by 60%",
        "• Conducted deep-dive analysis on healthcare quality metrics and patient satisfaction scores",
        "• Partnered with IT to develop data governance frameworks ensuring HIPAA compliance",
        "• Led cross-functional teams in revenue cycle optimization projects"
    ]
    for bullet in bullets2:
        doc.add_paragraph(bullet)

    doc.add_paragraph()

    # Education
    doc.add_heading("EDUCATION", level=2)
    edu = doc.add_paragraph()
    edu.add_run("Master of Business Administration (MBA)").bold = True
    edu.add_run(" - Healthcare Management Focus\n")
    edu.add_run("University of Louisville | Louisville, KY | 2019\n\n")
    edu.add_run("Bachelor of Science in Business Analytics").bold = True
    edu.add_run("\nUniversity of Kentucky | Lexington, KY | 2017")

    doc.add_paragraph()

    # Certifications
    doc.add_heading("CERTIFICATIONS & TECHNICAL SKILLS", level=2)
    certs = doc.add_paragraph()
    certs.add_run("• Certified Healthcare Data Analyst (CHDA)\n")
    certs.add_run("• SQL, Python, R, SAS\n")
    certs.add_run("• Tableau, Power BI, Looker\n")
    certs.add_run("• Epic, Cerner, Allscripts EMR systems\n")
    certs.add_run("• Advanced Excel, Statistical Analysis, Machine Learning")

    # Save the resume
    output_path = "documents/resumes/tailored/Matthew_Scott_Resume_Cigna_Nov2025.docx"
    doc.save(output_path)
    print(f"✅ Resume created: {output_path}")
    return output_path


def create_cigna_cover_letter():
    """Create personalized cover letter for Cigna position"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))

    doc.add_paragraph()

    # Recipient
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("The Cigna Group")
    doc.add_paragraph("Remote Position - Business Analytics Senior Advisor")

    doc.add_paragraph()

    # Salutation
    doc.add_paragraph("Dear Hiring Manager,")

    doc.add_paragraph()

    # Opening paragraph
    opening = doc.add_paragraph()
    opening.add_run(
        "I am writing to express my strong interest in the Business Analytics Senior Advisor position "
        "at The Cigna Group. With over 5 years of experience in healthcare analytics and a proven track "
        "record of driving operational improvements through data-driven insights, I am confident I would "
        "be a valuable addition to your analytics team. Your focus on using analytics to improve healthcare "
        "outcomes aligns perfectly with my professional passion and expertise."
    )

    doc.add_paragraph()

    # Body paragraph 1 - Match experience
    body1 = doc.add_paragraph()
    body1.add_run(
        "My experience directly aligns with your requirements for this role. I have spent the past five "
        "years analyzing and forecasting healthcare data for operational and business planning, with a "
        "particular focus on payer operations and clinical workflows. In my current role, I've led "
        "multiple business process improvement initiatives that resulted in significant cost savings and "
        "efficiency gains, including a 30% reduction in claims processing time and a 25% improvement in "
        "forecast accuracy for healthcare utilization."
    )

    doc.add_paragraph()

    # Body paragraph 2 - Specific achievements
    body2 = doc.add_paragraph()
    body2.add_run(
        "I bring deep knowledge of healthcare terminology and workflows, gained through hands-on experience "
        "with major health systems and payers. Key achievements that demonstrate my readiness for this role include:\n"
    )

    # Add bullet points
    doc.add_paragraph(
        "• Developed predictive models that reduced patient readmission rates by 18% across multiple hospitals",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Created executive dashboards and reporting systems that improved decision-making speed by 40%",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Led cross-functional teams in implementing data governance frameworks ensuring regulatory compliance",
        style='List Bullet'
    )

    doc.add_paragraph()

    # Body paragraph 3 - Why Cigna
    body3 = doc.add_paragraph()
    body3.add_run(
        "I am particularly drawn to Cigna's commitment to making healthcare simpler, more affordable, and "
        "more accessible. Your innovative approach to using analytics to improve member experiences and health "
        "outcomes resonates with my professional values. I am excited about the opportunity to contribute to "
        "Cigna's mission while leveraging your comprehensive benefits and professional development programs "
        "to continue growing my expertise."
    )

    doc.add_paragraph()

    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run(
        "I am excited about the opportunity to bring my healthcare analytics expertise to The Cigna Group "
        "and contribute to your continued success. I would welcome the chance to discuss how my background "
        "in healthcare operations, business process improvement, and data analytics can benefit your team. "
        "Thank you for considering my application."
    )

    doc.add_paragraph()

    # Sign-off
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("Matthew Scott")

    # Save the cover letter
    output_path = "documents/cover_letters/sent/Matthew_Scott_CoverLetter_Cigna_Nov2025.docx"
    doc.save(output_path)
    print(f"✅ Cover letter created: {output_path}")
    return output_path


def main():
    """Create complete application package for Cigna"""
    print("\n🎯 Creating Application Package for Cigna Business Analytics Senior Advisor")
    print("=" * 70)

    # Create resume
    print("\n📄 Creating tailored resume...")
    resume_path = create_cigna_resume()

    # Create cover letter
    print("\n📝 Creating personalized cover letter...")
    cover_path = create_cigna_cover_letter()

    print("\n" + "=" * 70)
    print("✨ APPLICATION PACKAGE COMPLETE!")
    print("\n📋 Next Steps:")
    print("1. Review the documents for any personal adjustments")
    print("2. Visit: https://jobs.thecignagroup.com/us/en/c/technology-jobs")
    print("3. Search for 'Business Analytics Senior Advisor'")
    print("4. Submit application with these documents")
    print("5. Update the job status in the database")
    print("\n💰 Salary Range: $109,500 - $182,500")
    print("📍 Location: Remote")
    print("\nGood luck! 🍀")


if __name__ == "__main__":
    main()