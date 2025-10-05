#!/usr/bin/env python3
"""
Create a professional Word document resume from the corrected text file.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set document margins (1 inch all around)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============= HEADER =============
# Name
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('MATTHEW DAVID SCOTT')
name_run.font.name = 'Calibri'
name_run.font.size = Pt(20)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(0, 0, 0)

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('Louisville, KY | 502-345-0525 | matthewdscott7@gmail.com | linkedin.com/in/mscott77/')
contact_run.font.name = 'Calibri'
contact_run.font.size = Pt(11)

# Add spacing
doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SENIOR BUSINESS ANALYST | HEALTHCARE ANALYTICS & QUALITY EVALUATION SPECIALIST')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(12)
title_run.font.bold = True

# Tagline
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_run = tagline.add_run('Data-Driven Decision Making | Medicare/CMS Compliance | SQL & API Testing | Stakeholder Leadership')
tagline_run.font.name = 'Calibri'
tagline_run.font.size = Pt(11)
tagline_run.font.italic = True

# Add spacing
doc.add_paragraph()

# Professional Summary
summary = doc.add_paragraph()
summary_run = summary.add_run('Strategic business analyst with 15+ years driving data-informed decision making across healthcare compliance, quality evaluation, and stakeholder leadership. Proven expertise in Medicare/CMS regulatory compliance, SQL-based data validation, analytics-driven testing, and cross-functional project management at Fortune 50 scale.')
summary_run.font.name = 'Calibri'
summary_run.font.size = Pt(11)

# Add spacing
doc.add_paragraph()

# ============= PROFESSIONAL EXPERIENCE =============
exp_heading = doc.add_paragraph()
exp_heading_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
exp_heading_run.font.name = 'Calibri'
exp_heading_run.font.size = Pt(13)
exp_heading_run.font.bold = True
exp_heading_run.font.color.rgb = RGBColor(0, 0, 0)

# Add bottom border
exp_heading.paragraph_format.space_after = Pt(6)

# --- HUMANA ---
job1_title = doc.add_paragraph()
job1_title_run = job1_title.add_run('SENIOR RISK MANAGEMENT PROFESSIONAL II, BUSINESS ANALYST')
job1_title_run.font.name = 'Calibri'
job1_title_run.font.size = Pt(11)
job1_title_run.font.bold = True

job1_details = doc.add_paragraph()
job1_details_run = job1_details.add_run('HUMANA, INC., Louisville, KY (Hybrid/Remote)')
job1_details_run.font.name = 'Calibri'
job1_details_run.font.size = Pt(11)
job1_details_run.font.bold = True

job1_dates = doc.add_paragraph()
job1_dates_run = job1_dates.add_run('November 2022 – August 2025')
job1_dates_run.font.name = 'Calibri'
job1_dates_run.font.size = Pt(11)
job1_dates_run.font.italic = True

job1_desc = doc.add_paragraph()
job1_desc_run = job1_desc.add_run('Medicare Compliance & Quality Assurance Business Analyst leading regulatory compliance, data validation, and stakeholder coordination for Humana\'s Medicare platform serving millions of consumers. Subject matter expert (SME) for CMS regulations with responsibility for cross-functional project leadership, SQL-based data analysis, and executive reporting.')
job1_desc_run.font.name = 'Calibri'
job1_desc_run.font.size = Pt(11)

# Career Progression
prog_heading = doc.add_paragraph()
prog_heading_run = prog_heading.add_run('Career Progression:')
prog_heading_run.font.name = 'Calibri'
prog_heading_run.font.size = Pt(11)
prog_heading_run.font.italic = True
prog_heading.paragraph_format.space_before = Pt(6)

progressions = [
    'Senior Risk Management Professional, May 2021 – Oct 2022',
    'Risk Management Professional 2, Sep 2017 – Apr 2021',
    'Risk Management Analyst, Jan 2016 – Aug 2017'
]

for prog in progressions:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(prog)
    p_run.font.name = 'Calibri'
    p_run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

# Key Achievements
ach_heading = doc.add_paragraph()
ach_heading_run = ach_heading.add_run('Key Achievements:')
ach_heading_run.font.name = 'Calibri'
ach_heading_run.font.size = Pt(11)
ach_heading_run.font.italic = True
ach_heading.paragraph_format.space_before = Pt(6)

achievements_humana = [
    'Led e-Commerce Acceleration and Data Modernization projects with ZERO critical or high defects through comprehensive testing and stakeholder coordination',
    'Orchestrated Annual Enrollment Period (AEP) web marketing preparation and execution—organization\'s highest-revenue period—ensuring seamless launch and CMS compliance',
    'Achieved 100% on-time delivery for 100+ time-sensitive Medicare Member Annual Communications (MAC) PDFs for 3 consecutive years as DRI',
    'Reviewed 400+ page CMS Final Rule guidance annually to identify potential risk areas and strategic implications, translating regulatory requirements into actionable business strategies',
    'Delivered executive-level monthly scorecards to senior leadership detailing team performance, trend analysis, and compliance metrics, driving data-informed strategic decisions',
    'Developed analytics-based testing methodology targeting highest-impact customer segments, creating test experiments that continually refined quality processes',
    'Built and maintained trusted relationships across IT, Legal, Regulatory & Compliance, and business units at enterprise scale, coordinating cross-functional engagements and managing third-party audit processes'
]

for ach in achievements_humana:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(ach)
    p_run.font.name = 'Calibri'
    p_run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

# Add spacing before next job
doc.add_paragraph()

# --- MIGHTILY ---
job2_title = doc.add_paragraph()
job2_title_run = job2_title.add_run('ACCOUNT EXECUTIVE')
job2_title_run.font.name = 'Calibri'
job2_title_run.font.size = Pt(11)
job2_title_run.font.bold = True

job2_details = doc.add_paragraph()
job2_details_run = job2_details.add_run('MIGHTILY, Louisville, KY')
job2_details_run.font.name = 'Calibri'
job2_details_run.font.size = Pt(11)
job2_details_run.font.bold = True

job2_dates = doc.add_paragraph()
job2_dates_run = job2_dates.add_run('July 2015 – December 2016')
job2_dates_run.font.name = 'Calibri'
job2_dates_run.font.size = Pt(11)
job2_dates_run.font.italic = True

job2_desc = doc.add_paragraph()
job2_desc_run = job2_desc.add_run('Client Relationship Manager and Strategic Account Lead managing 10+ digital marketing retainer accounts representing $40k in monthly recurring revenue. Served as intermediary between client project needs and internal technical teams across diverse industries, leading project delivery and revenue growth initiatives. Promoted from Rich Content Developer (2012-2015).')
job2_desc_run.font.name = 'Calibri'
job2_desc_run.font.size = Pt(11)

# Key Achievements
ach2_heading = doc.add_paragraph()
ach2_heading_run = ach2_heading.add_run('Key Achievements:')
ach2_heading_run.font.name = 'Calibri'
ach2_heading_run.font.size = Pt(11)
ach2_heading_run.font.italic = True
ach2_heading.paragraph_format.space_before = Pt(6)

achievements_mightily = [
    'Managed 10+ digital marketing retainer accounts accounting for $40k in monthly revenue with metric-based reporting and strategic recommendations',
    'Secured $50k website deal through consultative selling and managed primary account executive responsibilities for contracts ranging from $500k to $3M',
    'Created 60 Google Map location pins for Trilogy Health Services, significantly improving web presence and local market reach',
    'Launched Google AdWords display campaigns using A/B testing to maximize ad spend efficiency and ROI'
]

for ach in achievements_mightily:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(ach)
    p_run.font.name = 'Calibri'
    p_run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

# Add spacing
doc.add_paragraph()

# ============= EDUCATION =============
edu_heading = doc.add_paragraph()
edu_heading_run = edu_heading.add_run('EDUCATION & CERTIFICATION')
edu_heading_run.font.name = 'Calibri'
edu_heading_run.font.size = Pt(13)
edu_heading_run.font.bold = True
edu_heading.paragraph_format.space_after = Pt(6)

edu1 = doc.add_paragraph()
edu1_run = edu1.add_run('BACHELOR OF SCIENCE IN COMMUNICATION | University of Louisville, Expected 2025')
edu1_run.font.name = 'Calibri'
edu1_run.font.size = Pt(11)
edu1_run.font.bold = True

edu2 = doc.add_paragraph()
edu2_run = edu2.add_run('Dean\'s List, Fall 2013')
edu2_run.font.name = 'Calibri'
edu2_run.font.size = Pt(11)

doc.add_paragraph()

cert = doc.add_paragraph()
cert_run = cert.add_run('Google Ads Certified')
cert_run.font.name = 'Calibri'
cert_run.font.size = Pt(11)

# Add spacing
doc.add_paragraph()

# ============= SKILLS =============
skills_heading = doc.add_paragraph()
skills_heading_run = skills_heading.add_run('CORE COMPETENCIES & TECHNICAL SKILLS')
skills_heading_run.font.name = 'Calibri'
skills_heading_run.font.size = Pt(13)
skills_heading_run.font.bold = True
skills_heading.paragraph_format.space_after = Pt(6)

# Healthcare & Compliance
hc_heading = doc.add_paragraph()
hc_heading_run = hc_heading.add_run('HEALTHCARE & COMPLIANCE')
hc_heading_run.font.name = 'Calibri'
hc_heading_run.font.size = Pt(11)
hc_heading_run.font.bold = True

hc_content = doc.add_paragraph()
hc_content_run = hc_content.add_run('Medicare/CMS Regulations | Regulatory Compliance | Risk Assessment | Quality Assurance | ESP (Enterprise Solution Point) | EIP (Enterprise Integration Platform) | ETDM (Enterprise Test Data Management) | Healthcare Domain Expertise')
hc_content_run.font.name = 'Calibri'
hc_content_run.font.size = Pt(11)

# Data & Analytics
da_heading = doc.add_paragraph()
da_heading_run = da_heading.add_run('DATA & ANALYTICS')
da_heading_run.font.name = 'Calibri'
da_heading_run.font.size = Pt(11)
da_heading_run.font.bold = True
da_heading.paragraph_format.space_before = Pt(6)

da_content = doc.add_paragraph()
da_content_run = da_content.add_run('SQL (MS SQL Server, MySQL) | API Testing & Validation | Analytics-Based Testing | A/B Testing & Optimization | Data Validation | Trend Analysis | Performance Metrics')
da_content_run.font.name = 'Calibri'
da_content_run.font.size = Pt(11)

# Leadership & PM
lpm_heading = doc.add_paragraph()
lpm_heading_run = lpm_heading.add_run('LEADERSHIP & PROJECT MANAGEMENT')
lpm_heading_run.font.name = 'Calibri'
lpm_heading_run.font.size = Pt(11)
lpm_heading_run.font.bold = True
lpm_heading.paragraph_format.space_before = Pt(6)

lpm_content = doc.add_paragraph()
lpm_content_run = lpm_content.add_run('Stakeholder Engagement | Cross-Functional Leadership | Agile/SDLC | Sprint Planning | User Stories & UAT (User Acceptance Testing) | Executive Reporting | Vendor Management | Team Mentoring')
lpm_content_run.font.name = 'Calibri'
lpm_content_run.font.size = Pt(11)

# Technical Tools
tt_heading = doc.add_paragraph()
tt_heading_run = tt_heading.add_run('TECHNICAL TOOLS & PLATFORMS')
tt_heading_run.font.name = 'Calibri'
tt_heading_run.font.size = Pt(11)
tt_heading_run.font.bold = True
tt_heading.paragraph_format.space_before = Pt(6)

tt_content = doc.add_paragraph()
tt_content_run = tt_content.add_run('JIRA | SharePoint | MS Teams | Slack | Oracle | Braze | Firebase | Splunk | Applitools | Beyond Compare | WordPress | Power Apps | Power Automate')
tt_content_run.font.name = 'Calibri'
tt_content_run.font.size = Pt(11)

# Save document
doc.save('/Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume.docx')

print("SUCCESS: Resume created at /Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume.docx")
