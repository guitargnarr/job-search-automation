#!/usr/bin/env python3
"""
Create a professionally styled Word document resume inspired by Adobe templates.
Two-column layout with modern design elements.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph, width=100, color='000000'):
    """Add a horizontal line below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Size in eighths of a point
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# Create document
doc = Document()

# Set narrow margins for two-column layout
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# ============= HEADER (Full Width) =============
# Name
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('MATTHEW DAVID SCOTT')
name_run.font.name = 'Calibri'
name_run.font.size = Pt(24)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(31, 78, 120)  # Professional blue
name.paragraph_format.space_after = Pt(2)

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('Louisville, KY  |  502-345-0525  |  matthewdscott7@gmail.com  |  linkedin.com/in/mscott77/')
contact_run.font.name = 'Calibri'
contact_run.font.size = Pt(10)
contact.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SENIOR BUSINESS ANALYST  |  HEALTHCARE ANALYTICS & QUALITY EVALUATION SPECIALIST')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(11)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 78, 120)
title.paragraph_format.space_after = Pt(8)
add_horizontal_line(title, color='1F4E78')

# Professional Summary
summary_heading = doc.add_paragraph()
summary_heading_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
summary_heading_run.font.name = 'Calibri'
summary_heading_run.font.size = Pt(11)
summary_heading_run.font.bold = True
summary_heading_run.font.color.rgb = RGBColor(31, 78, 120)
summary_heading.paragraph_format.space_before = Pt(8)
summary_heading.paragraph_format.space_after = Pt(4)

summary = doc.add_paragraph()
summary_run = summary.add_run('Strategic business analyst with 15+ years driving data-informed decision making across healthcare compliance, quality evaluation, and stakeholder leadership. Proven expertise in Medicare/CMS regulatory compliance, SQL-based data validation, analytics-driven testing, and cross-functional project management at Fortune 50 scale.')
summary_run.font.name = 'Calibri'
summary_run.font.size = Pt(10)
summary.paragraph_format.space_after = Pt(10)

# Create table for two-column layout
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Set column widths (left wider for experience)
left_cell = table.rows[0].cells[0]
right_cell = table.rows[0].cells[1]

left_cell.width = Inches(4.2)
right_cell.width = Inches(2.3)

# Remove cell borders
for cell in table.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ============= LEFT COLUMN - EXPERIENCE =============
left_paragraphs = left_cell.paragraphs[0]
left_paragraphs.text = ''

def add_to_left(text='', size=10, bold=False, italic=False, color_rgb=None, space_before=0, space_after=0):
    p = left_cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# PROFESSIONAL EXPERIENCE
exp_heading = add_to_left('PROFESSIONAL EXPERIENCE', 12, bold=True, color_rgb=RGBColor(31, 78, 120), space_after=6)
add_horizontal_line(exp_heading, color='1F4E78')

# Humana
add_to_left('SENIOR RISK MANAGEMENT PROFESSIONAL II', 10, bold=True, space_before=6)
add_to_left('BUSINESS ANALYST', 10, bold=True)
add_to_left('HUMANA, INC.  |  Louisville, KY (Hybrid/Remote)', 9, italic=True, color_rgb=RGBColor(89, 89, 89))
add_to_left('November 2022 – August 2025', 9, italic=True, color_rgb=RGBColor(89, 89, 89), space_after=4)

add_to_left('Medicare Compliance & Quality Assurance Business Analyst leading regulatory compliance, data validation, and stakeholder coordination for Humana\'s Medicare platform serving millions of consumers. Subject matter expert (SME) for CMS regulations with responsibility for cross-functional project leadership, SQL-based data analysis, and executive reporting.', 10, space_after=6)

# Career Progression
add_to_left('Career Progression:', 9, bold=True, italic=True)
prog_items = [
    'Senior Risk Management Professional, May 2021 – Oct 2022',
    'Risk Management Professional 2, Sep 2017 – Apr 2021',
    'Risk Management Analyst, Jan 2016 – Aug 2017'
]
for item in prog_items:
    p = left_cell.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)

# Key Achievements
add_to_left('Key Achievements:', 9, bold=True, italic=True, space_before=4)
achievements = [
    'Led e-Commerce Acceleration and Data Modernization projects with ZERO critical or high defects',
    'Orchestrated Annual Enrollment Period (AEP)—organization\'s highest-revenue period',
    'Achieved 100% on-time delivery for 100+ Medicare PDFs for 3 consecutive years as DRI',
    'Reviewed 400+ page CMS Final Rule guidance annually to identify strategic risk areas',
    'Delivered executive-level monthly scorecards to senior leadership',
    'Developed analytics-based testing methodology targeting highest-impact customer segments',
    'Built trusted relationships across IT, Legal, R&C at enterprise scale'
]
for ach in achievements:
    p = left_cell.add_paragraph(style='List Bullet')
    run = p.add_run(ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)

# Mightily
add_to_left('ACCOUNT EXECUTIVE', 10, bold=True, space_before=10)
add_to_left('MIGHTILY  |  Louisville, KY', 9, italic=True, color_rgb=RGBColor(89, 89, 89))
add_to_left('July 2015 – December 2016', 9, italic=True, color_rgb=RGBColor(89, 89, 89), space_after=4)

add_to_left('Client Relationship Manager managing 10+ digital marketing accounts representing $40k in monthly recurring revenue. Led project delivery and revenue growth initiatives across diverse industries. Promoted from Rich Content Developer (2012-2015).', 10, space_after=6)

add_to_left('Key Achievements:', 9, bold=True, italic=True)
mightily_ach = [
    'Managed 10+ accounts accounting for $40k monthly revenue',
    'Secured $50k website deal; managed contracts ranging from $500k to $3M',
    'Created 60 Google Map location pins for Trilogy Health Services',
    'Launched Google AdWords campaigns using A/B testing for ROI optimization'
]
for ach in mightily_ach:
    p = left_cell.add_paragraph(style='List Bullet')
    run = p.add_run(ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)

# ============= RIGHT COLUMN - SKILLS & EDUCATION =============
right_paragraphs = right_cell.paragraphs[0]
right_paragraphs.text = ''

def add_to_right(text='', size=10, bold=False, italic=False, color_rgb=None, space_before=0, space_after=0):
    p = right_cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# EDUCATION
edu_heading = add_to_right('EDUCATION', 11, bold=True, color_rgb=RGBColor(31, 78, 120), space_after=4)
add_horizontal_line(edu_heading, color='1F4E78')

add_to_right('Bachelor of Science', 9, bold=True, space_before=4)
add_to_right('Communication', 9, bold=True)
add_to_right('University of Louisville', 9)
add_to_right('Expected 2025', 9, italic=True, color_rgb=RGBColor(89, 89, 89), space_after=2)
add_to_right('Dean\'s List, Fall 2013', 8, space_after=10)

# CERTIFICATION
cert_heading = add_to_right('CERTIFICATION', 11, bold=True, color_rgb=RGBColor(31, 78, 120), space_after=4)
add_horizontal_line(cert_heading, color='1F4E78')
add_to_right('Google Ads Certified', 9, space_before=4, space_after=10)

# CORE COMPETENCIES
comp_heading = add_to_right('CORE COMPETENCIES', 11, bold=True, color_rgb=RGBColor(31, 78, 120), space_after=4)
add_horizontal_line(comp_heading, color='1F4E78')

add_to_right('Healthcare & Compliance', 9, bold=True, space_before=4, space_after=2)
comp_hc = [
    'Medicare/CMS Regulations',
    'Risk Assessment',
    'Quality Assurance',
    'ESP/EIP/ETDM',
    'Healthcare Domain'
]
for item in comp_hc:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + item)
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(1)

add_to_right('Data & Analytics', 9, bold=True, space_before=6, space_after=2)
comp_da = [
    'SQL (MS SQL, MySQL)',
    'API Testing',
    'Analytics Testing',
    'A/B Testing',
    'Trend Analysis'
]
for item in comp_da:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + item)
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(1)

add_to_right('Leadership & PM', 9, bold=True, space_before=6, space_after=2)
comp_lpm = [
    'Stakeholder Engagement',
    'Cross-Functional Lead',
    'Agile/SDLC',
    'Sprint Planning',
    'UAT & User Stories',
    'Executive Reporting',
    'Vendor Management'
]
for item in comp_lpm:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + item)
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(1)

add_to_right('Technical Tools', 9, bold=True, space_before=6, space_after=2)
comp_tools = [
    'JIRA',
    'SharePoint',
    'MS Teams',
    'Oracle',
    'Braze',
    'Firebase',
    'Splunk',
    'Applitools',
    'Power Apps/Automate'
]
for item in comp_tools:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + item)
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(1)

# Save document
doc.save('/Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume_Styled.docx')

print("SUCCESS: Styled resume created at /Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume_Styled.docx")
