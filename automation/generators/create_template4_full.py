#!/usr/bin/env python3
"""
Template 4 Style: Two-column professional (no photo)
Inspired by William Jones template - modern two-column
FULL CONTENT VERSION - Under 2 pages
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color='F2F2F2'):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Create two-column table
table = doc.add_table(rows=1, cols=2)
table.autofit = False
left_cell = table.rows[0].cells[0]
right_cell = table.rows[0].cells[1]
left_cell.width = Inches(4.5)
right_cell.width = Inches(2.5)

# Shade right column
shade_cell(right_cell, 'D9D9D9')

# Remove borders
for cell in table.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# LEFT COLUMN
left_cell.paragraphs[0].text = ''

def add_left(text='', size=9, bold=False, italic=False, color=None, space_before=0, space_after=0):
    p = left_cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# NAME
add_left('MATTHEW DAVID', 20, bold=True, space_after=0)
add_left('SCOTT', 20, bold=True, space_after=4)

# TITLE
add_left('SENIOR BUSINESS ANALYST', 10, bold=True, space_after=8)

# ABOUT ME
add_left('PROFESSIONAL SUMMARY', 11, bold=True, space_after=4)
add_left('Strategic business analyst with 15+ years driving data-informed decision making across healthcare compliance, quality evaluation, and stakeholder leadership. Proven expertise in Medicare/CMS regulatory compliance, SQL-based data validation, analytics-driven testing, and cross-functional project management at Fortune 50 scale.', 8, space_after=8)

# EXPERIENCE
add_left('PROFESSIONAL EXPERIENCE', 11, bold=True, space_after=4)

# Humana
add_left('HUMANA, INC.', 9, bold=True, space_after=1)
add_left('Senior Risk Management Professional II, Business Analyst', 8, italic=True, space_after=1)
add_left('November 2022 – August 2025', 8, color=RGBColor(128, 128, 128), space_after=3)

add_left('Medicare Compliance & Quality Assurance Business Analyst leading regulatory compliance, data validation, and stakeholder coordination for Humana\'s Medicare platform serving millions of consumers. Subject matter expert (SME) for CMS regulations with responsibility for cross-functional project leadership, SQL-based data analysis, and executive reporting.', 8, space_after=3)

add_left('Career Progression:', 8, bold=True, italic=True, space_after=2)
progs = ['Senior Risk Management Professional, May 2021 – Oct 2022', 'Risk Management Professional 2, Sep 2017 – Apr 2021', 'Risk Management Analyst, Jan 2016 – Aug 2017']
for prog in progs:
    p = left_cell.add_paragraph()
    run = p.add_run('• ' + prog)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

add_left('Key Achievements:', 8, bold=True, italic=True, space_before=2, space_after=2)
achs = [
    'Led e-Commerce Acceleration and Data Modernization projects with ZERO critical or high defects through comprehensive testing and stakeholder coordination',
    'Orchestrated Annual Enrollment Period (AEP) web marketing preparation and execution—organization\'s highest-revenue period—ensuring seamless launch and CMS compliance',
    'Achieved 100% on-time delivery for 100+ time-sensitive Medicare Member Annual Communications (MAC) PDFs for 3 consecutive years as DRI',
    'Reviewed 400+ page CMS Final Rule guidance annually to identify potential risk areas and strategic implications, translating regulatory requirements into actionable business strategies',
    'Delivered executive-level monthly scorecards to senior leadership detailing team performance, trend analysis, and compliance metrics, driving data-informed strategic decisions',
    'Developed analytics-based testing methodology targeting highest-impact customer segments, creating test experiments that continually refined quality processes',
    'Built and maintained trusted relationships across IT, Legal, Regulatory & Compliance, and business units at enterprise scale, coordinating cross-functional engagements and managing third-party audit processes'
]
for ach in achs:
    p = left_cell.add_paragraph()
    run = p.add_run('• ' + ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

# Mightily
add_left('MIGHTILY', 9, bold=True, space_before=6, space_after=1)
add_left('Account Executive', 8, italic=True, space_after=1)
add_left('July 2015 – December 2016', 8, color=RGBColor(128, 128, 128), space_after=3)

add_left('Client Relationship Manager and Strategic Account Lead managing 10+ digital marketing retainer accounts representing $40k in monthly recurring revenue. Served as intermediary between client project needs and internal technical teams across diverse industries, leading project delivery and revenue growth initiatives. Promoted from Rich Content Developer (2012-2015).', 8, space_after=3)

add_left('Key Achievements:', 8, bold=True, italic=True, space_after=2)
m_achs = [
    'Managed 10+ digital marketing retainer accounts accounting for $40k in monthly revenue with metric-based reporting and strategic recommendations',
    'Secured $50k website deal through consultative selling and managed primary account executive responsibilities for contracts ranging from $500k to $3M',
    'Created 60 Google Map location pins for Trilogy Health Services, significantly improving web presence and local market reach',
    'Launched Google AdWords display campaigns using A/B testing to maximize ad spend efficiency and ROI'
]
for ach in m_achs:
    p = left_cell.add_paragraph()
    run = p.add_run('• ' + ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

# RIGHT COLUMN
right_cell.paragraphs[0].text = ''

def add_right(text='', size=9, bold=False, italic=False, color=None, space_before=0, space_after=0):
    p = right_cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# CONTACT
add_right('CONTACT', 11, bold=True, space_after=4)
add_right('Louisville, KY', 8, space_after=2)
add_right('502-345-0525', 8, space_after=2)
add_right('matthewdscott7@gmail.com', 7, space_after=2)
add_right('linkedin.com/in/mscott77/', 7, space_after=10)

# EDUCATION
add_right('EDUCATION', 11, bold=True, space_after=4)
add_right('Bachelor of Science', 8, bold=True, space_after=1)
add_right('Communication', 8, space_after=1)
add_right('University of Louisville', 8, space_after=1)
add_right('Expected 2025', 8, italic=True, space_after=2)
add_right('Dean\'s List, Fall 2013', 7, space_after=10)

# CERTIFICATION
add_right('CERTIFICATION', 11, bold=True, space_after=4)
add_right('Google Ads Certified', 8, space_after=10)

# SKILLS
add_right('CORE COMPETENCIES', 11, bold=True, space_after=4)

add_right('Healthcare & Compliance', 8, bold=True, space_after=2)
hc = ['Medicare/CMS Regulations', 'Regulatory Compliance', 'Risk Assessment', 'Quality Assurance', 'ESP/EIP/ETDM']
for skill in hc:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + skill)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

add_right('Data & Analytics', 8, bold=True, space_before=4, space_after=2)
da = ['SQL', 'API Testing', 'Analytics Testing', 'A/B Testing', 'Trend Analysis']
for skill in da:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + skill)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

add_right('Leadership & PM', 8, bold=True, space_before=4, space_after=2)
lpm = ['Stakeholder Engagement', 'Cross-Functional Leadership', 'Agile/SDLC', 'Executive Reporting', 'Team Mentoring']
for skill in lpm:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + skill)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

add_right('Technical Tools', 8, bold=True, space_before=4, space_after=2)
tools = ['JIRA', 'SharePoint', 'MS Teams', 'Oracle', 'Splunk', 'Power Apps']
for tool in tools:
    p = right_cell.add_paragraph()
    run = p.add_run('• ' + tool)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.space_after = Pt(1)

doc.save('/Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume_Template4_Full.docx')
print("SUCCESS: Template 4 created")
