#!/usr/bin/env python3
"""
Template 2 Style: Single-column with decorative header
Inspired by MARY SMITH template - professional with star borders
FULL CONTENT VERSION - Under 2 pages
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph, color='000000'):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# NAME
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('MATTHEW DAVID SCOTT')
name_run.font.name = 'Calibri'
name_run.font.size = Pt(22)
name_run.font.bold = True
name.paragraph_format.space_after = Pt(2)

# TITLE with borders
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SENIOR BUSINESS ANALYST | HEALTHCARE ANALYTICS & QUALITY EVALUATION SPECIALIST')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(10)
title_run.font.bold = True
title.paragraph_format.space_after = Pt(4)
add_horizontal_line(title)

# Contact
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('502-345-0525 | matthewdscott7@gmail.com | Louisville, KY | linkedin.com/in/mscott77/')
contact_run.font.name = 'Calibri'
contact_run.font.size = Pt(9)
contact.paragraph_format.space_after = Pt(8)

# Summary
summary = doc.add_paragraph()
summary_run = summary.add_run('Strategic business analyst with 15+ years driving data-informed decision making across healthcare compliance, quality evaluation, and stakeholder leadership. Proven expertise in Medicare/CMS regulatory compliance, SQL-based data validation, analytics-driven testing, and cross-functional project management at Fortune 50 scale.')
summary_run.font.name = 'Calibri'
summary_run.font.size = Pt(9)
summary.paragraph_format.space_after = Pt(8)

# EXPERIENCE
exp_h = doc.add_paragraph()
exp_h_run = exp_h.add_run('PROFESSIONAL EXPERIENCE')
exp_h_run.font.name = 'Calibri'
exp_h_run.font.size = Pt(12)
exp_h_run.font.bold = True
exp_h.paragraph_format.space_after = Pt(4)

# Humana
job1 = doc.add_paragraph()
job1_run = job1.add_run('Senior Risk Management Professional II, Business Analyst | HUMANA, INC. | Nov 2022 – Aug 2025')
job1_run.font.name = 'Calibri'
job1_run.font.size = Pt(9)
job1_run.font.bold = True
job1.paragraph_format.space_after = Pt(2)

desc1 = doc.add_paragraph()
desc1_run = desc1.add_run('Medicare Compliance & Quality Assurance Business Analyst leading regulatory compliance, data validation, and stakeholder coordination for Humana\'s Medicare platform serving millions of consumers. Subject matter expert (SME) for CMS regulations with responsibility for cross-functional project leadership, SQL-based data analysis, and executive reporting.')
desc1_run.font.name = 'Calibri'
desc1_run.font.size = Pt(8)
desc1.paragraph_format.space_after = Pt(3)

# Career progression
cp = doc.add_paragraph()
cp_run = cp.add_run('Career Progression:')
cp_run.font.name = 'Calibri'
cp_run.font.size = Pt(8)
cp_run.font.italic = True
cp_run.font.bold = True
cp.paragraph_format.space_after = Pt(2)

progs = ['Senior Risk Management Professional, May 2021 – Oct 2022', 'Risk Management Professional 2, Sep 2017 – Apr 2021', 'Risk Management Analyst, Jan 2016 – Aug 2017']
for prog in progs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(prog)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(1)

# Key achievements
ka = doc.add_paragraph()
ka_run = ka.add_run('Key Achievements:')
ka_run.font.name = 'Calibri'
ka_run.font.size = Pt(8)
ka_run.font.italic = True
ka_run.font.bold = True
ka.paragraph_format.space_before = Pt(2)
ka.paragraph_format.space_after = Pt(2)

achs = [
    'Led e-Commerce Acceleration and Data Modernization projects with ZERO critical or high defects through comprehensive testing and stakeholder coordination',
    'Orchestrated Annual Enrollment Period (AEP) web marketing preparation and execution—organization\'s highest-revenue period—ensuring seamless launch and CMS compliance',
    'Achieved 100% on-time delivery for 100+ time-sensitive Medicare Member Annual Communications (MAC) PDFs for 3 consecutive years as DRI',
    'Reviewed 400+ page CMS Final Rule guidance annually to identify potential risk areas and strategic implications, translating regulatory requirements into actionable business strategies',
    'Delivered executive-level monthly scorecards to senior leadership detailing team performance, trend analysis, and compliance metrics, driving data-informed strategic decisions',
    'Developed analytics-based testing methodology targeting highest-impact customer segments, creating test experiments that continually refined quality processes',
    'Built and maintained trusted relationships across IT, Legal, Regulatory & Compliance, and business units at enterprise scale, coordinating cross-functional engagements and managing third-party audit processes'
]
for ach in achs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(1)

# Mightily
job2 = doc.add_paragraph()
job2_run = job2.add_run('Account Executive | MIGHTILY | July 2015 – December 2016')
job2_run.font.name = 'Calibri'
job2_run.font.size = Pt(9)
job2_run.font.bold = True
job2.paragraph_format.space_before = Pt(6)
job2.paragraph_format.space_after = Pt(2)

desc2 = doc.add_paragraph()
desc2_run = desc2.add_run('Client Relationship Manager and Strategic Account Lead managing 10+ digital marketing retainer accounts representing $40k in monthly recurring revenue. Served as intermediary between client project needs and internal technical teams across diverse industries, leading project delivery and revenue growth initiatives. Promoted from Rich Content Developer (2012-2015).')
desc2_run.font.name = 'Calibri'
desc2_run.font.size = Pt(8)
desc2.paragraph_format.space_after = Pt(3)

ka2 = doc.add_paragraph()
ka2_run = ka2.add_run('Key Achievements:')
ka2_run.font.name = 'Calibri'
ka2_run.font.size = Pt(8)
ka2_run.font.italic = True
ka2_run.font.bold = True
ka2.paragraph_format.space_after = Pt(2)

m_achs = [
    'Managed 10+ digital marketing retainer accounts accounting for $40k in monthly revenue with metric-based reporting and strategic recommendations',
    'Secured $50k website deal through consultative selling and managed primary account executive responsibilities for contracts ranging from $500k to $3M',
    'Created 60 Google Map location pins for Trilogy Health Services, significantly improving web presence and local market reach',
    'Launched Google AdWords display campaigns using A/B testing to maximize ad spend efficiency and ROI'
]
for ach in m_achs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ach)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(1)

# EDUCATION
edu_h = doc.add_paragraph()
edu_h_run = edu_h.add_run('EDUCATION & CERTIFICATION')
edu_h_run.font.name = 'Calibri'
edu_h_run.font.size = Pt(12)
edu_h_run.font.bold = True
edu_h.paragraph_format.space_before = Pt(8)
edu_h.paragraph_format.space_after = Pt(4)

edu = doc.add_paragraph()
edu_run = edu.add_run('Bachelor of Science in Communication | University of Louisville, Expected 2025')
edu_run.font.name = 'Calibri'
edu_run.font.size = Pt(9)
edu_run.font.bold = True
edu.paragraph_format.space_after = Pt(1)

dl = doc.add_paragraph()
dl_run = dl.add_run('Dean\'s List, Fall 2013')
dl_run.font.name = 'Calibri'
dl_run.font.size = Pt(8)
dl.paragraph_format.space_after = Pt(4)

cert = doc.add_paragraph()
cert_run = cert.add_run('Google Ads Certified')
cert_run.font.name = 'Calibri'
cert_run.font.size = Pt(9)
cert.paragraph_format.space_after = Pt(8)

# SKILLS
skills_h = doc.add_paragraph()
skills_h_run = skills_h.add_run('CORE COMPETENCIES & TECHNICAL SKILLS')
skills_h_run.font.name = 'Calibri'
skills_h_run.font.size = Pt(12)
skills_h_run.font.bold = True
skills_h.paragraph_format.space_after = Pt(4)

# Healthcare
hc_h = doc.add_paragraph()
hc_h_run = hc_h.add_run('Healthcare & Compliance: ')
hc_h_run.font.name = 'Calibri'
hc_h_run.font.size = Pt(8)
hc_h_run.font.bold = True
hc_c_run = hc_h.add_run('Medicare/CMS Regulations | Regulatory Compliance | Risk Assessment | Quality Assurance | ESP (Enterprise Solution Point) | EIP (Enterprise Integration Platform) | ETDM (Enterprise Test Data Management) | Healthcare Domain Expertise')
hc_c_run.font.name = 'Calibri'
hc_c_run.font.size = Pt(8)
hc_h.paragraph_format.space_after = Pt(3)

# Data
da_h = doc.add_paragraph()
da_h_run = da_h.add_run('Data & Analytics: ')
da_h_run.font.name = 'Calibri'
da_h_run.font.size = Pt(8)
da_h_run.font.bold = True
da_c_run = da_h.add_run('SQL (MS SQL Server, MySQL) | API Testing & Validation | Analytics-Based Testing | A/B Testing & Optimization | Data Validation | Trend Analysis | Performance Metrics')
da_c_run.font.name = 'Calibri'
da_c_run.font.size = Pt(8)
da_h.paragraph_format.space_after = Pt(3)

# Leadership
lpm_h = doc.add_paragraph()
lpm_h_run = lpm_h.add_run('Leadership & Project Management: ')
lpm_h_run.font.name = 'Calibri'
lpm_h_run.font.size = Pt(8)
lpm_h_run.font.bold = True
lpm_c_run = lpm_h.add_run('Stakeholder Engagement | Cross-Functional Leadership | Agile/SDLC | Sprint Planning | User Stories & UAT (User Acceptance Testing) | Executive Reporting | Vendor Management | Team Mentoring')
lpm_c_run.font.name = 'Calibri'
lpm_c_run.font.size = Pt(8)
lpm_h.paragraph_format.space_after = Pt(3)

# Tools
tools_h = doc.add_paragraph()
tools_h_run = tools_h.add_run('Technical Tools & Platforms: ')
tools_h_run.font.name = 'Calibri'
tools_h_run.font.size = Pt(8)
tools_h_run.font.bold = True
tools_c_run = tools_h.add_run('JIRA | SharePoint | MS Teams | Slack | Oracle | Braze | Firebase | Splunk | Applitools | Beyond Compare | WordPress | Power Apps | Power Automate')
tools_c_run.font.name = 'Calibri'
tools_c_run.font.size = Pt(8)

doc.save('/Users/matthewscott/Desktop/Job_Search/Matthew_Scott_Resume_Template2_Full.docx')
print("SUCCESS: Template 2 created")
