#!/usr/bin/env python3
"""
Create 3 tiers of cover letter templates as Word documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cover_letter(filename, tier_name, content_blocks):
    """Create a cover letter Word document"""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    header = doc.add_paragraph()
    header_run = header.add_run('MATTHEW DAVID SCOTT')
    header_run.font.name = 'Calibri'
    header_run.font.size = Pt(14)
    header_run.font.bold = True

    contact = doc.add_paragraph()
    contact_run = contact.add_run('Louisville, KY  |  502-345-0525  |  matthewdscott7@gmail.com  |  linkedin.com/in/mscott77/')
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(10)
    contact.paragraph_format.space_after = Pt(12)

    # Date placeholder
    date = doc.add_paragraph()
    date_run = date.add_run('[DATE]')
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(11)
    date.paragraph_format.space_after = Pt(12)

    # Recipient info
    recipient = doc.add_paragraph()
    recipient_run = recipient.add_run('[HIRING MANAGER NAME]\n[TITLE]\n[COMPANY NAME]\n[COMPANY ADDRESS]')
    recipient_run.font.name = 'Calibri'
    recipient_run.font.size = Pt(11)
    recipient.paragraph_format.space_after = Pt(12)

    # Salutation
    salutation = doc.add_paragraph()
    salutation_run = salutation.add_run('Dear [HIRING MANAGER NAME/Hiring Manager],')
    salutation_run.font.name = 'Calibri'
    salutation_run.font.size = Pt(11)
    salutation.paragraph_format.space_after = Pt(12)

    # Body paragraphs
    for block in content_blocks:
        p = doc.add_paragraph()
        p_run = p.add_run(block)
        p_run.font.name = 'Calibri'
        p_run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # Closing
    closing = doc.add_paragraph()
    closing_run = closing.add_run('Thank you for considering my application. I look forward to speaking with you soon.')
    closing_run.font.name = 'Calibri'
    closing_run.font.size = Pt(11)
    closing.paragraph_format.space_after = Pt(12)

    signature = doc.add_paragraph()
    signature_run = signature.add_run('Sincerely,\n\nMatthew David Scott')
    signature_run.font.name = 'Calibri'
    signature_run.font.size = Pt(11)

    # Save
    doc.save(filename)

# ==================== TIER 1: HEALTHCARE SPECIALIST ====================
tier1_blocks = [
    'I am writing to express my strong interest in the [JOB TITLE] position at [COMPANY NAME]. With over 9 years of specialized experience in healthcare compliance, regulatory analysis, and data-driven quality assurance at Fortune 50 scale, I am confident I can deliver immediate value to your team while advancing [COMPANY\'S] mission in [HEALTHCARE DOMAIN/MISSION].',

    'In my current role as Senior Risk Management Professional II at Humana, I serve as the Medicare web space compliance subject matter expert, where I have successfully led mission-critical initiatives that directly align with the requirements of this position:\n\n• Led e-Commerce Acceleration and Data Modernization projects with zero critical defects through comprehensive cross-functional stakeholder coordination across IT, Legal, and Regulatory & Compliance teams\n• Orchestrated Annual Enrollment Period (AEP) web marketing preparation and execution—Humana\'s highest-revenue period—ensuring seamless launch while maintaining 100% CMS compliance\n• Achieved 100% on-time delivery for 100+ time-sensitive Medicare Member Annual Communications (MAC) PDFs for three consecutive years as the Directly Responsible Individual (DRI)\n• Reviewed 400+ page CMS Final Rule guidance annually to identify potential risk areas and strategic implications, translating complex regulatory requirements into actionable business strategies',

    'My expertise in Medicare/CMS regulations, combined with proven proficiency in SQL-based data validation, analytics-driven testing methodologies, and executive-level reporting, positions me uniquely to contribute to [COMPANY\'S SPECIFIC INITIATIVE/TEAM]. I am particularly drawn to [COMPANY NAME] because [SPECIFIC REASON - company mission, healthcare focus, technology innovation, recent news, etc.].',

    'Beyond technical capabilities, my track record demonstrates excellence in building trusted relationships and coordinating complex initiatives at enterprise scale. I have successfully partnered with stakeholders across multiple departments to deliver strategic outcomes, manage third-party audit processes, and provide data-informed insights that drive executive decision-making.',

    'I am excited about the opportunity to bring my healthcare domain expertise, regulatory compliance acumen, and passion for quality-driven innovation to [COMPANY NAME]. I would welcome the chance to discuss how my background aligns with your team\'s goals and how I can contribute to [SPECIFIC COMPANY OBJECTIVE].'
]

create_cover_letter('/Users/matthewscott/Desktop/Job_Search/Cover_Letter_Tier1_Healthcare.docx', 'Tier 1', tier1_blocks)
print("✅ Tier 1 Cover Letter Created (Healthcare Specialist)")

# ==================== TIER 2: BUSINESS ANALYST GENERALIST ====================
tier2_blocks = [
    'I am writing to apply for the [JOB TITLE] position at [COMPANY NAME]. With 15+ years of experience driving data-informed decision making, cross-functional project leadership, and analytical innovation, I am enthusiastic about contributing to [COMPANY\'S] continued success.',

    'Throughout my career, I have demonstrated a consistent ability to translate complex data and regulatory requirements into actionable business strategies. At Humana, I led high-stakes projects including e-Commerce Acceleration and Data Modernization initiatives that launched with zero critical defects through comprehensive testing protocols and stakeholder coordination. My analytical testing methodology, which targets highest-impact customer segments through data-driven experimentation, has continually refined quality processes and improved business outcomes.',

    'Key achievements that demonstrate my capabilities:\n\n• Delivered executive-level monthly scorecards to senior leadership detailing team performance, trend analysis, and compliance metrics, driving strategic decisions\n• Built and maintained trusted relationships across IT, Legal, Regulatory & Compliance, and business units at enterprise scale, managing third-party audit processes\n• Managed 10+ digital marketing retainer accounts accounting for $40k in monthly recurring revenue with metric-based reporting\n• Developed comprehensive risk assessments and systematic validation protocols across multiple business-critical scenarios',

    'My proficiency in SQL (MS SQL Server, MySQL), API testing, Agile/SDLC methodologies, and executive reporting—combined with experience managing projects from initiation through delivery—aligns well with the requirements for this role. I am particularly impressed by [COMPANY\'S SPECIFIC ACCOMPLISHMENT/MISSION/INNOVATION] and am excited about the opportunity to contribute to [SPECIFIC INITIATIVE/TEAM].',

    'I bring a collaborative approach to problem-solving, strong communication skills for both technical and business audiences, and a proven track record of delivering results in fast-paced, high-stakes environments. I would welcome the opportunity to discuss how my analytical expertise and cross-functional leadership can support [COMPANY NAME]\'s objectives.'
]

create_cover_letter('/Users/matthewscott/Desktop/Job_Search/Cover_Letter_Tier2_GeneralBA.docx', 'Tier 2', tier2_blocks)
print("✅ Tier 2 Cover Letter Created (General Business Analyst)")

# ==================== TIER 3: AI TRANSITION ====================
tier3_blocks = [
    'I am writing to apply for the [JOB TITLE] position at [COMPANY NAME]. As a business analyst with extensive experience in quality evaluation, data validation, and analytical testing methodologies, I am eager to contribute to AI development and model evaluation while transitioning my expertise to the artificial intelligence domain.',

    'My background in quality assurance and data-driven decision making at Humana has prepared me exceptionally well for AI evaluation work:\n\n• Designed and executed comprehensive test cases to validate system outputs against expected performance criteria, analyzing results to detect patterns in errors and system failures—skills directly transferable to AI model evaluation\n• Developed analytics-based testing methodology that targets highest-impact scenarios through data-driven experimentation and continual refinement\n• Conducted systematic data validation using SQL across business-critical applications, ensuring accuracy and quality at scale\n• Produced detailed analysis reports for leadership, translating complex technical performance into clear, actionable insights',

    'I am particularly drawn to this opportunity because it combines my strengths in analytical evaluation, attention to detail, and pattern recognition with the exciting frontier of generative AI technology. My experience evaluating system performance, identifying edge cases, and providing detailed feedback to development teams translates naturally to assessing AI model outputs, identifying biases or errors, and improving model performance through quality annotation and prompt engineering.',

    'With a Bachelor of Science in Communication and professional expertise in both technical analysis and clear communication, I am well-positioned to evaluate AI responses for accuracy, coherence, and alignment with intended outcomes. I am passionate about contributing to AI development that is reliable, ethical, and beneficial, and I am committed to bringing the same rigor and attention to detail that has defined my career in healthcare quality assurance.',

    'I would be delighted to discuss how my analytical background, quality-focused mindset, and enthusiasm for AI technology can contribute to [COMPANY NAME]\'s work in advancing generative AI capabilities.'
]

create_cover_letter('/Users/matthewscott/Desktop/Job_Search/Cover_Letter_Tier3_AI_Transition.docx', 'Tier 3', tier3_blocks)
print("✅ Tier 3 Cover Letter Created (AI Transition)")

print("\n" + "="*60)
print("ALL 3 COVER LETTER TEMPLATES CREATED SUCCESSFULLY")
print("="*60)
