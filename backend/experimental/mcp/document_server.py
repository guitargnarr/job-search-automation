#!/usr/bin/env python3
"""
Document Generation MCP Server
Creates ATS-optimized resumes and personalized cover letters
"""

import os
import sys
import json
import asyncio
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path
import hashlib

# Add backend to path for imports
sys.path.append(str(Path(__file__).parent.parent.parent))

from mcp.server.fastmcp import FastMCP
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import existing services and models
from backend.core.database import get_session, engine, Base
from backend.models.models import Job, Company, Application
from backend.services.ats_optimizer import ATSOptimizer
from backend.core.logging import get_logger

logger = get_logger(__name__)

# Initialize MCP server
mcp = FastMCP("Document Generation")

# Initialize ATS optimizer
ats_optimizer = ATSOptimizer()


# ====================
# Resources (Data Access)
# ====================

@mcp.resource("templates://resumes")
async def get_resume_templates() -> str:
    """Get available resume templates and formats"""
    templates = {
        "ats_optimized": {
            "name": "ATS-Optimized Template",
            "description": "Clean, simple format optimized for ATS parsing",
            "sections": [
                "Contact Information",
                "Professional Summary",
                "Core Skills",
                "Professional Experience",
                "Education",
                "Certifications"
            ],
            "features": [
                "Simple formatting",
                "Standard section headers",
                "Bullet points for achievements",
                "Keywords prominently displayed",
                "No tables or columns",
                "Standard fonts (Arial, Calibri)"
            ]
        },
        "executive": {
            "name": "Executive Template",
            "description": "Professional format for senior positions",
            "sections": [
                "Contact Information",
                "Executive Summary",
                "Core Competencies",
                "Professional Experience",
                "Key Achievements",
                "Education & Credentials",
                "Board Positions & Affiliations"
            ],
            "features": [
                "Emphasis on leadership",
                "Quantified achievements",
                "Strategic accomplishments",
                "Industry recognition"
            ]
        },
        "technical": {
            "name": "Technical Template",
            "description": "Format for technical and IT roles",
            "sections": [
                "Contact Information",
                "Technical Summary",
                "Technical Skills",
                "Programming Languages & Tools",
                "Professional Experience",
                "Projects",
                "Education",
                "Certifications"
            ],
            "features": [
                "Detailed technical skills",
                "Project highlights",
                "Technology stack emphasis",
                "GitHub/portfolio links"
            ]
        }
    }

    return json.dumps(templates, indent=2)


@mcp.resource("templates://cover_letters")
async def get_cover_letter_templates() -> str:
    """Get cover letter templates for different scenarios"""
    templates = {
        "standard": {
            "name": "Standard Cover Letter",
            "structure": [
                "Opening: Reference specific position and how you learned about it",
                "Body Paragraph 1: Match your experience to key requirements",
                "Body Paragraph 2: Highlight relevant achievements with metrics",
                "Body Paragraph 3: Demonstrate knowledge of company and culture fit",
                "Closing: Express enthusiasm and next steps"
            ],
            "tone": "Professional and enthusiastic"
        },
        "career_change": {
            "name": "Career Change Cover Letter",
            "structure": [
                "Opening: Acknowledge career transition and enthusiasm for new field",
                "Body Paragraph 1: Transferable skills and their relevance",
                "Body Paragraph 2: Relevant training, certifications, or projects",
                "Body Paragraph 3: Why this change and why this company",
                "Closing: Commitment to the transition and value proposition"
            ],
            "tone": "Confident and purposeful"
        },
        "referral": {
            "name": "Referral Cover Letter",
            "structure": [
                "Opening: Mention referral source immediately",
                "Body Paragraph 1: Connection's endorsement and your qualifications",
                "Body Paragraph 2: Specific value you bring to the role",
                "Body Paragraph 3: Alignment with company goals",
                "Closing: Thank referrer and express interest in discussion"
            ],
            "tone": "Warm and confident"
        }
    }

    return json.dumps(templates, indent=2)


@mcp.resource("keywords://extracted")
async def get_job_keywords(job_id: int) -> str:
    """Get extracted keywords from a specific job posting"""
    try:
        async with get_session() as session:
            result = await session.execute(
                select(Job).where(Job.id == job_id)
            )
            job = result.scalar_one_or_none()

            if not job or not job.job_description:
                return json.dumps({"error": f"Job {job_id} not found or has no description"})

            # Extract keywords using ATS optimizer
            keywords = ats_optimizer.extract_keywords(job.job_description)
            skills = ats_optimizer.extract_skills(job.job_description)
            requirements = ats_optimizer.extract_requirements(job.job_description)

            return json.dumps({
                "job_id": job_id,
                "company": job.company_name,
                "title": job.title,
                "keywords": {
                    "top_keywords": keywords[:20],
                    "technical_skills": skills.get("technical", []),
                    "soft_skills": skills.get("soft", []),
                    "requirements": requirements,
                    "industry_terms": skills.get("industry", [])
                }
            }, indent=2)

    except Exception as e:
        logger.error(f"Error extracting keywords: {str(e)}")
        return json.dumps({"error": str(e)})


@mcp.resource("documents://generated")
async def list_generated_documents() -> str:
    """List all generated documents with their metadata"""
    try:
        doc_dir = Path("documents")
        resume_dir = doc_dir / "resumes" / "tailored"
        cover_dir = doc_dir / "cover_letters" / "sent"

        documents = []

        # List tailored resumes
        if resume_dir.exists():
            for resume_file in resume_dir.glob("*.docx"):
                stat = resume_file.stat()
                documents.append({
                    "type": "resume",
                    "filename": resume_file.name,
                    "path": str(resume_file),
                    "size_kb": round(stat.st_size / 1024, 1),
                    "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })

        # List cover letters
        if cover_dir.exists():
            for cover_file in cover_dir.glob("*.docx"):
                stat = cover_file.stat()
                documents.append({
                    "type": "cover_letter",
                    "filename": cover_file.name,
                    "path": str(cover_file),
                    "size_kb": round(stat.st_size / 1024, 1),
                    "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })

        return json.dumps({
            "total_documents": len(documents),
            "documents": sorted(documents, key=lambda x: x["modified"], reverse=True)
        }, indent=2)

    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        return json.dumps({"error": str(e)})


# ====================
# Tools (Actions)
# ====================

@mcp.tool()
async def generate_tailored_resume(
    job_id: int,
    base_resume_path: str,
    template: str = "ats_optimized",
    emphasis_areas: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    Generate an ATS-optimized resume tailored for a specific job

    Args:
        job_id: ID of the job to tailor for
        base_resume_path: Path to the base resume file
        template: Template to use (ats_optimized, executive, technical)
        emphasis_areas: Areas to emphasize (e.g., ["leadership", "analytics"])

    Returns:
        Dictionary with generated resume path and optimization score
    """
    try:
        async with get_session() as session:
            # Get job details
            result = await session.execute(
                select(Job).where(Job.id == job_id)
            )
            job = result.scalar_one_or_none()

            if not job:
                return {"success": False, "error": f"Job {job_id} not found"}

            if not job.job_description:
                return {"success": False, "error": "Job has no description to optimize against"}

            # Check if base resume exists
            base_path = Path(base_resume_path)
            if not base_path.exists():
                return {"success": False, "error": f"Base resume not found: {base_resume_path}"}

            # Extract keywords from job description
            keywords = ats_optimizer.extract_keywords(job.job_description)
            skills = ats_optimizer.extract_skills(job.job_description)

            # Optimize resume content
            optimization_result = ats_optimizer.optimize_resume(
                base_resume_path,
                job.job_description,
                job.title
            )

            # Create output directory
            output_dir = Path("documents/resumes/tailored")
            output_dir.mkdir(parents=True, exist_ok=True)

            # Generate filename
            company_safe = job.company_name.replace(" ", "_").replace("/", "_")
            timestamp = datetime.now().strftime("%Y%m%d")
            output_filename = f"resume_{company_safe}_{job_id}_{timestamp}.docx"
            output_path = output_dir / output_filename

            # Create the tailored document
            doc = Document()

            # Configure document settings for ATS
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Add content (simplified for demonstration)
            # In production, this would parse and optimize the base resume
            doc.add_heading('MATTHEW SCOTT', 0)

            # Contact info
            contact = doc.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.add_run('Louisville, KY | matthewdscott7@gmail.com | LinkedIn')

            # Professional Summary with keywords
            doc.add_heading('PROFESSIONAL SUMMARY', 1)
            summary = doc.add_paragraph()
            summary_text = f"Results-driven {job.title} with expertise in {', '.join(keywords[:5])}. "
            summary_text += f"Proven track record in {', '.join(skills.get('technical', [])[:3])}."
            summary.add_run(summary_text)

            # Core Skills section with job keywords
            doc.add_heading('CORE SKILLS', 1)
            skills_para = doc.add_paragraph()
            all_skills = keywords[:15] + skills.get('technical', [])[:10]
            skills_para.add_run(' • '.join(all_skills))

            # Save the document
            doc.save(str(output_path))

            # Calculate final ATS score
            with open(output_path, 'rb') as f:
                content = f.read()
                # Simple scoring based on keyword presence
                score = min(100, 70 + len([k for k in keywords if k.lower() in str(content).lower()]))

            return {
                "success": True,
                "job_id": job_id,
                "company": job.company_name,
                "title": job.title,
                "output_path": str(output_path),
                "filename": output_filename,
                "ats_score": score,
                "keywords_included": len([k for k in keywords if k.lower() in str(content).lower()]),
                "optimization": {
                    "keywords_matched": optimization_result.get("keywords_matched", 0),
                    "missing_keywords": optimization_result.get("missing_keywords", [])[:5],
                    "score_improvement": optimization_result.get("score_improvement", 0)
                }
            }

    except Exception as e:
        logger.error(f"Error generating tailored resume: {str(e)}")
        return {"success": False, "error": str(e)}


@mcp.tool()
async def create_cover_letter(
    job_id: int,
    tone: str = "professional",
    highlights: Optional[List[str]] = None,
    referral_name: Optional[str] = None
) -> Dict[str, Any]:
    """
    Create a personalized cover letter for a specific job

    Args:
        job_id: ID of the job to apply for
        tone: Tone of the letter (professional, enthusiastic, confident)
        highlights: Specific achievements to highlight
        referral_name: Name of person who referred you (if applicable)

    Returns:
        Dictionary with generated cover letter path
    """
    try:
        async with get_session() as session:
            # Get job and company details
            result = await session.execute(
                select(Job).where(Job.id == job_id)
            )
            job = result.scalar_one_or_none()

            if not job:
                return {"success": False, "error": f"Job {job_id} not found"}

            # Get company information
            company_result = await session.execute(
                select(Company).where(Company.name == job.company_name)
            )
            company = company_result.scalar_one_or_none()

            # Create output directory
            output_dir = Path("documents/cover_letters/sent")
            output_dir.mkdir(parents=True, exist_ok=True)

            # Generate filename
            company_safe = job.company_name.replace(" ", "_").replace("/", "_")
            timestamp = datetime.now().strftime("%Y%m%d")
            output_filename = f"cover_letter_{company_safe}_{job_id}_{timestamp}.docx"
            output_path = output_dir / output_filename

            # Create the cover letter document
            doc = Document()

            # Header with date
            date_para = doc.add_paragraph()
            date_para.add_run(datetime.now().strftime("%B %d, %Y"))

            # Hiring manager address
            doc.add_paragraph()
            doc.add_paragraph("Hiring Manager")
            doc.add_paragraph(job.company_name)
            if job.location:
                doc.add_paragraph(job.location)

            doc.add_paragraph()

            # Salutation
            if referral_name:
                doc.add_paragraph(f"Dear Hiring Team (Referred by {referral_name}),")
            else:
                doc.add_paragraph("Dear Hiring Manager,")

            doc.add_paragraph()

            # Opening paragraph
            opening = doc.add_paragraph()
            if referral_name:
                opening_text = f"I am writing to express my strong interest in the {job.title} position at {job.company_name}, "
                opening_text += f"as recommended by {referral_name}. "
            else:
                opening_text = f"I am writing to express my strong interest in the {job.title} position at {job.company_name}. "

            opening_text += "With my background in healthcare analytics and business intelligence, "
            opening_text += "I am confident I would be a valuable addition to your team."
            opening.add_run(opening_text)

            doc.add_paragraph()

            # Body paragraph 1 - Match experience to requirements
            body1 = doc.add_paragraph()
            body1_text = f"My experience aligns well with the requirements for this {job.title} role. "
            if highlights:
                body1_text += f"Specifically, I have {highlights[0]}. "
            body1_text += "Throughout my career, I have consistently delivered data-driven insights "
            body1_text += "that have improved operational efficiency and strategic decision-making."
            body1.add_run(body1_text)

            doc.add_paragraph()

            # Body paragraph 2 - Specific achievements
            body2 = doc.add_paragraph()
            body2_text = "In my previous roles, I have successfully:"
            body2.add_run(body2_text)

            # Add bullet points for achievements
            if highlights and len(highlights) > 1:
                for highlight in highlights[1:4]:
                    doc.add_paragraph(f"• {highlight}", style='List Bullet')
            else:
                doc.add_paragraph("• Developed automated reporting systems reducing manual effort by 60%", style='List Bullet')
                doc.add_paragraph("• Led cross-functional teams in implementing data governance frameworks", style='List Bullet')
                doc.add_paragraph("• Delivered actionable insights resulting in $2M+ cost savings", style='List Bullet')

            doc.add_paragraph()

            # Body paragraph 3 - Company fit
            body3 = doc.add_paragraph()
            if company and company.culture_notes:
                body3_text = f"I am particularly drawn to {job.company_name} because of {company.culture_notes[:100]}. "
            else:
                body3_text = f"I am particularly drawn to {job.company_name}'s commitment to innovation in healthcare. "
            body3_text += "Your focus on data-driven decision making aligns perfectly with my professional philosophy "
            body3_text += "and expertise."
            body3.add_run(body3_text)

            doc.add_paragraph()

            # Closing paragraph
            closing = doc.add_paragraph()
            closing_text = f"I am excited about the opportunity to contribute to {job.company_name}'s continued success. "
            closing_text += "I would welcome the chance to discuss how my skills and experience can benefit your team. "
            closing_text += "Thank you for considering my application."
            closing.add_run(closing_text)

            doc.add_paragraph()

            # Sign-off
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph("Matthew Scott")

            # Save the document
            doc.save(str(output_path))

            return {
                "success": True,
                "job_id": job_id,
                "company": job.company_name,
                "title": job.title,
                "output_path": str(output_path),
                "filename": output_filename,
                "tone": tone,
                "referral_included": bool(referral_name),
                "highlights_included": len(highlights) if highlights else 0
            }

    except Exception as e:
        logger.error(f"Error creating cover letter: {str(e)}")
        return {"success": False, "error": str(e)}


@mcp.tool()
async def analyze_ats_score(
    resume_path: str,
    job_description: str
) -> Dict[str, Any]:
    """
    Analyze a resume's ATS compatibility score against a job description

    Args:
        resume_path: Path to the resume file
        job_description: Job description text

    Returns:
        Dictionary with ATS score and recommendations
    """
    try:
        # Check if resume exists
        resume_file = Path(resume_path)
        if not resume_file.exists():
            return {"success": False, "error": f"Resume not found: {resume_path}"}

        # Calculate ATS score
        score = ats_optimizer.calculate_ats_score(resume_path, job_description)

        # Extract keywords from job description
        job_keywords = ats_optimizer.extract_keywords(job_description)

        # Analyze resume content
        resume_analysis = ats_optimizer.analyze_resume(resume_path)

        # Find missing keywords
        resume_keywords = resume_analysis.get("keywords", [])
        missing_keywords = [k for k in job_keywords if k.lower() not in [r.lower() for r in resume_keywords]]

        # Generate recommendations
        recommendations = []
        if score < 70:
            recommendations.append("Add more relevant keywords from the job description")
        if len(missing_keywords) > 10:
            recommendations.append(f"Include these missing keywords: {', '.join(missing_keywords[:5])}")
        if resume_analysis.get("format_issues"):
            recommendations.extend(resume_analysis["format_issues"])

        return {
            "success": True,
            "ats_score": score,
            "score_interpretation": "Excellent" if score >= 85 else "Good" if score >= 70 else "Needs Improvement",
            "analysis": {
                "keywords_found": len([k for k in job_keywords if k.lower() in [r.lower() for r in resume_keywords]]),
                "total_job_keywords": len(job_keywords),
                "missing_critical_keywords": missing_keywords[:10],
                "format_score": resume_analysis.get("format_score", 0)
            },
            "recommendations": recommendations,
            "estimated_pass_rate": "High" if score >= 80 else "Medium" if score >= 65 else "Low"
        }

    except Exception as e:
        logger.error(f"Error analyzing ATS score: {str(e)}")
        return {"success": False, "error": str(e)}


@mcp.tool()
async def extract_requirements(
    job_url: Optional[str] = None,
    job_description: Optional[str] = None
) -> Dict[str, Any]:
    """
    Extract requirements from a job posting

    Args:
        job_url: URL of the job posting (optional)
        job_description: Job description text (optional)

    Returns:
        Dictionary with extracted requirements and qualifications
    """
    try:
        if not job_description and not job_url:
            return {"success": False, "error": "Provide either job_url or job_description"}

        # If URL provided, note that actual scraping would be done here
        if job_url and not job_description:
            return {
                "success": False,
                "error": "URL scraping not implemented. Please provide job_description text."
            }

        # Extract requirements using ATS optimizer
        requirements = ats_optimizer.extract_requirements(job_description)
        skills = ats_optimizer.extract_skills(job_description)

        # Parse into categories
        result = {
            "success": True,
            "requirements": {
                "education": [],
                "experience": [],
                "technical_skills": skills.get("technical", []),
                "soft_skills": skills.get("soft", []),
                "certifications": [],
                "other": []
            }
        }

        # Categorize requirements (simplified)
        for req in requirements:
            req_lower = req.lower()
            if any(word in req_lower for word in ["degree", "bachelor", "master", "phd"]):
                result["requirements"]["education"].append(req)
            elif any(word in req_lower for word in ["years", "experience", "experienced"]):
                result["requirements"]["experience"].append(req)
            elif any(word in req_lower for word in ["certification", "certified", "license"]):
                result["requirements"]["certifications"].append(req)
            else:
                result["requirements"]["other"].append(req)

        # Add summary
        result["summary"] = {
            "total_requirements": len(requirements),
            "technical_skills_count": len(skills.get("technical", [])),
            "soft_skills_count": len(skills.get("soft", [])),
            "prioritization": "Focus on technical skills and experience requirements"
        }

        return result

    except Exception as e:
        logger.error(f"Error extracting requirements: {str(e)}")
        return {"success": False, "error": str(e)}


# ====================
# Prompts (Templates)
# ====================

@mcp.prompt("application_package")
async def application_package_prompt() -> str:
    """Complete application package generation"""
    return """Create a complete application package for a job:

1. Analyze the job posting:
   - Extract all requirements and keywords
   - Identify company culture and values
   - Note specific skills and experiences emphasized

2. Generate tailored resume:
   - Use ATS-optimized template
   - Include all relevant keywords naturally
   - Quantify achievements where possible
   - Ensure 80+ ATS score

3. Create personalized cover letter:
   - Reference specific job requirements
   - Highlight 3-4 relevant achievements
   - Show knowledge of company
   - Express genuine enthusiasm

4. Prepare supplementary documents:
   - Portfolio pieces if relevant
   - Reference list if requested
   - Work samples if applicable

5. Quality check:
   - Verify ATS score > 80
   - Check for typos and formatting
   - Ensure consistency across documents
   - Confirm all requirements addressed"""


async def main():
    """Run the MCP server"""
    # Initialize database
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    logger.info("Document Generation MCP Server starting...")
    await mcp.run()


if __name__ == "__main__":
    asyncio.run(main())